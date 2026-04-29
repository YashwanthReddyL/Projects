from docx import Document

doc = Document()

# ================= TITLE PAGE =================
doc.add_heading('ML/DL Training in HPC Environment', 0)
doc.add_paragraph('Comparison of Single System vs HPC Cluster')
doc.add_paragraph(' ')
doc.add_paragraph('Prepared for: Academic / HPC Study Report')
doc.add_paragraph('Topic: High Performance Computing in Machine Learning and Deep Learning')
doc.add_paragraph(' ')

doc.add_page_break()

# ================= INTRO =================
doc.add_heading('Introduction', level=1)
doc.add_paragraph(
    "This document compares Machine Learning and Deep Learning training on a single system "
    "versus an HPC cluster environment. The focus is on performance, scalability, limitations, "
    "and system behavior in high-performance computing contexts."
)

# ================= ADVANTAGES TABLE =================
doc.add_heading('Advantages', level=1)

table1 = doc.add_table(rows=1, cols=2)
hdr = table1.rows[0].cells
hdr[0].text = "System"
hdr[1].text = "Advantages"

data1 = [
("Single System",
 "Simple setup; low cost; easy debugging; no network dependency; low overhead; good for small models"),

("HPC Cluster",
 "High compute power; distributed training; scalability; fault tolerance; parallel execution; efficient resource usage; supports large-scale ML/DL")
]

for item in data1:
    row = table1.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]

# ================= DISADVANTAGES TABLE =================
doc.add_heading('Disadvantages', level=1)

table2 = doc.add_table(rows=1, cols=2)
hdr = table2.rows[0].cells
hdr[0].text = "System"
hdr[1].text = "Disadvantages"

data2 = [
("Single System",
 "Limited compute power; GPU/VRAM bottleneck; no scalability; slow training; no fault tolerance; dataset limitations; no distributed training"),

("HPC Cluster",
 "High cost; complex setup; network dependency; communication overhead; difficult debugging; scheduling delays; high maintenance")
]

for item in data2:
    row = table2.add_row().cells
    row[0].text = item[0]
    row[1].text = item[1]

# ================= CONCLUSION =================
doc.add_heading('Conclusion', level=1)
doc.add_paragraph(
    "A single system is suitable for small-scale and experimental ML/DL workloads due to its simplicity "
    "and low cost, but it is limited by hardware constraints. In contrast, HPC clusters provide scalable "
    "and high-performance computing capabilities essential for large-scale deep learning and scientific "
    "computations, despite their higher complexity and cost."
)

# ================= SAVE =================
doc.save("ML_DL_HPC_Report.docx")

print("Document created: ML_DL_HPC_Report.docx")